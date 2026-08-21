# -*- coding: utf-8 -*-
"""Generate QA deliverables: Design-and-Editorial-QA-Report.docx and
Terminology-and-Consistency-Check.xlsx."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, '..', 'qa')
os.makedirs(OUT, exist_ok=True)

PURPLE = '7C32C9'; NAVY = '2A206A'; TEAL = '00ABAF'; INK = '232046'

# ---------------------------------------------------------------- DOCX helpers
def set_rtl(par):
    pPr = par._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def ar(doc, text, size=11, bold=False, color=INK, style_h=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    r = p.add_run(text)
    r.font.name = 'Diodrum Arabic'
    r._element.rPr.rFonts.set(qn('w:cs'), 'Diodrum Arabic')
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    return p

def h1(doc, text):  return ar(doc, text, 16, True, PURPLE)
def h2(doc, text):  return ar(doc, text, 13, True, TEAL)
def body(doc, text): return ar(doc, text, 10.5, False, '474956')

def ar_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # RTL table
    tblPr = t._tbl.tblPr
    bidi = OxmlElement('w:bidiVisual'); tblPr.append(bidi)
    for j, htxt in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p)
        r = p.add_run(htxt); r.font.bold = True; r.font.size = Pt(9.5)
        r.font.name = 'Diodrum Arabic'
        r.font.color.rgb = RGBColor.from_string('FFFFFF')
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), NAVY)
        c._tc.get_or_add_tcPr().append(shd)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[1 + i].cells[j]
            c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p)
            r = p.add_run(str(val)); r.font.size = Pt(9.5)
            r.font.name = 'Diodrum Arabic'
    return t

# ---------------------------------------------------------------- QA report
doc = Document()
for section in doc.sections:
    section.left_margin = section.right_margin = Cm(2)

ar(doc, 'هيئة الحكومة الرقمية — برنامج «قدراتك»', 11, True, TEAL)
ar(doc, 'تقرير الجودة التصميمية والتحريرية', 20, True, PURPLE)
ar(doc, 'Design and Editorial QA Report — الدراسة المعيارية لتطوير القيادات الحكومية العليا في الذكاء الاصطناعي (الإصدار 3.1 → الإخراج التنفيذي النهائي)', 11)
ar(doc, 'تاريخ التقرير: 21 أغسطس 2026', 10, False, '6E7077')

h1(doc, 'أولاً — ملخص ما تم إنجازه')
body(doc, 'حُوّلت النسخة v3.1 (ملفا DOCX العربي والإنجليزي) إلى إصدار تنفيذي مصمم وفق الهوية البصرية المعتمدة في ملف «التصميم بالألوان المعتمدة من ريان المصمم»، في نسختين متطابقتين (عربية وإنجليزية) من 39 صفحة لكل منهما، بصيغتي PPTX قابلة للتحرير بالكامل وPDF عالي الجودة بروابط حية. حُفظ الهيكل السداسي المعتمد مع قسم مستقل للمراجع [1]–[71]، وطُبق نظام ألوان المستويات الأربعة ضمن اللوحة المعتمدة (بنفسجي 7C32C9 للوزراء، كحلي 2A206A لنواب الوزراء وأصحاب المعالي، فيروزي 00ABAF للوكلاء، أخضر 1CC182 للوكلاء المساعدين).')
ar_table(doc,
 ['البند', 'القيمة'],
 [['عدد الصفحات (لكل نسخة)', '39 صفحة A4 عمودي'],
  ['الخط', 'Diodrum Arabic (جميع الأوزان، وفق هوية ريان — مطابق لخط ملف التصميم المرجعي)'],
  ['المراجع', '71 مرجعاً محفوظة نصاً وروابط، مع وسوم أدلة ملونة (S/E/P/R/N/I/J/T)'],
  ['الروابط الحية في PDF', '132 رابطاً قابلاً للنقر في كل نسخة'],
  ['التحقق التقني للملفات', 'اجتاز كلا ملفي PPTX فحص validate.py (بنية OOXML سليمة)'],
  ['العناصر البصرية', 'كل الجداول والبطاقات والرسوم عناصر أصلية قابلة للتحرير؛ لا صفحات مصورة'],
 ])

h1(doc, 'ثانياً — ما تم تعديله وتصحيحه')
h2(doc, '1) تصحيحات المصطلحات القيادية (النسخة الإنجليزية)')
body(doc, 'أزيلت جميع الاستخدامات المخالفة وفق القاموس الملزم، ونتيجة المسح الآلي النهائي: صفر ظهور لكلمة Undersecretary/Assistant Undersecretary في كامل النسخة الإنجليزية.')
ar_table(doc,
 ['الموضع في v3.1', 'قبل', 'بعد'],
 [['§2.4 العنوان والفئة', 'Level 3: Undersecretaries / Ministry undersecretaries', 'Level 3: Deputy Ministers, Deputy Governors, and Equivalent Leaders / deputy ministers'],
  ['§2.5 العنوان والفئة', 'Assistant Undersecretaries', 'Assistant Deputy Ministers'],
  ['§2.3 طبيعة الفئة', 'merged with undersecretaries', 'merged with deputy ministers'],
  ['§2.6 و§4.5 و§5.1 الجداول', 'Ministry undersecretary / Assistant undersecretary', 'Deputy minister / Assistant deputy minister'],
  ['§4.5 قراءة الجدول', 'ministry undersecretaries in the UAE executive programme', 'deputy ministers in the UAE executive programme'],
  ['§6.2 التوصية 3', 'between minister and undersecretary', 'between the minister and the deputy minister'],
  ['§2.3 العنوان', 'Excellency-rank Leaders', 'Vice Ministers and Excellency-Rank Leaders (الصيغة المعتمدة)'],
 ])
h2(doc, '2) تصحيحات نحوية عربية خفيفة (لا تمس المعنى)')
ar_table(doc,
 ['الموضع', 'قبل', 'بعد'],
 [['§3.3 كازاخستان', 'فيلق مسماة … مسماة ومساءل', 'فيلق مسمّى … مسمّى ومُساءل'],
  ['§4.3 العنوان', 'الوكلاء ونواب المحافظون', 'وكلاء الوزارات ونواب المحافظين'],
  ['§5.2', 'ومسار الوكلاء المساعدون', 'ومسار الوكلاء المساعدين'],
  ['§5.7', 'ولا يتراكم بقاءه', 'ولا يتراكم عند المالك (إعادة صياغة سليمة)'],
  ['§6.1 (3)', 'الطلب القيادي العليا قائم', 'الطلب لدى القيادات العليا قائم'],
 ])
h2(doc, '3) ما تم اختصاره أو إعادة توزيعه (دون حذف معلومة جوهرية)')
ar_table(doc,
 ['الموضع', 'الإجراء'],
 [['§1.4', 'وُزعت المؤشرات على صفحتين: بطاقات أرقام كبيرة للطلب العالمي، وبطاقتا مؤشر لموقع المملكة، مع صندوق «مؤشرات لا سببية»'],
  ['§2.2–2.5', 'قالب موحد لصفحات المستويات بالعناصر الستة الثابتة وبالترتيب نفسه: الفئة، الطبيعة، نوع القرار، البروتوكول، المخرج، قرار البناء/الشراء'],
  ['§3.2 (جدول 12 برنامجاً)', 'قُسم على صفحتين (6+6) مع صفحة قراءة مستقلة بثلاثة استنتاجات'],
  ['§3.3', 'جدول مضغوط لعشر دول + بطاقات حالات على صفحتين بشارات نوع الدليل (منفذ/هدف معلن/تقييم مستقل)'],
  ['§5.6', 'رقما التكلفة الاسترشادية أبرزا في بطاقتين كبيرتين مع صندوق حدود القراءة قبل جدول المكونات'],
  ['المراجع', 'أُخرجت في أربع صفحات بشارات أدلة ملونة وروابط على أسطر مستقلة قابلة للنقر'],
  ['التكرارات', 'حُفظت الفكرة في موضعها الأقوى (مثال: «هدف لا إنجاز» مفصلة في حالة الإمارات؛ التقييم البريطاني «غير الحاسم» مفصل في §3.3) مع إشارات مختصرة في المواضع الأخرى'],
 ])

h1(doc, 'ثالثاً — نتائج المراجعات الخمس الإلزامية')
ar_table(doc,
 ['المراجعة', 'النتيجة'],
 [['1. المصطلحات', 'اجتازت: صفر مصطلح محظور؛ Vice Minister (11)، Deputy Minister (25)، Assistant Deputy Minister (11)، Deputy Governor (6) في مواضعها الصحيحة؛ مساعد الوزير ثابت في المستوى الثاني في النسختين'],
  ['2. التطابق الثنائي', 'اجتازت: 39 = 39 صفحة؛ تسلسل الأقسام والجداول والرسوم والألوان متطابق؛ فحص آلي للأرقام صفحةً بصفحة أظهر تطابق القيم، والفروق المرصودة كلها اصطلاحية موثقة في ملف XLSX (أرقام مكتوبة كلمات بالعربية مثل «القسم الرابع»، وصيغة «45 ألفاً» مقابل 45,000)'],
  ['3. الأدلة', 'اجتازت: كل ادعاء مهم يحمل علامة مصدر؛ الإعلانات مقدمة إعلانات (الإمارات «هدف لا إنجاز»، سنغافورة «معلن ولم يبدأ»، لي كوان يو وبلافاتنك «مقررة وليست منفذة»)؛ التقديرات موسومة [65][69] بوسم T؛ التقييم البريطاني موصوف بغير الحاسم في كل المواضع؛ قيود الوصول مثبتة في عمود مستقل'],
  ['4. التصميم', 'اجتازت: التزام كامل بلوحة ريان وأسلوب الغلاف وفواصل الأقسام والبطاقات والجداول والتذييل؛ خلفية الغلاف وصورة المدينة مستخرجتان من ملف ريان نفسه؛ لا جداول مقصوصة ولا نصوص خارج الحدود (فحص بصري لكل الصفحات الـ78)'],
  ['5. السرد والصوت الواحد', 'اجتازت: انتقالات تحريرية من سطر أو سطرين في نهاية كل قسم تمهد للتالي؛ نبرة حكومية تنفيذية موحدة؛ عناوين استنتاجية لكل صفحة؛ قاعدة «يختلف البرنامج بحسب نوع القرار» تظهر في مواضعها المقررة فقط'],
 ])

h1(doc, 'رابعاً — نقاط تحتاج تحققاً بشرياً قبل الاعتماد النهائي')
body(doc, '1) قائمة المراجع في ملفي DOCX المصدرَين (العربي والإنجليزي) مكتوبة بصياغة عربية واحدة؛ وقد حُفظت حرفياً في النسختين حفاظاً على قابلية التحقق. يُقترح قرار بشري بشأن تعريب/تدويل أسماء الناشرين في النسخة الإنجليزية في إصدار لاحق.')
body(doc, '2) الرسوم والأسعار والتواريخ تخضع للتحوط القياسي المنصوص في الدراسة: تُراجع مع المزود قبل أي خطوة شرائية — وبخاصة سعر جونز هوبكنز [22] (سبق عرض نسخة أقصر بسعر أدنى)، وموعد الخصم المبكر لمدرسة لي كوان يو (30 سبتمبر 2026)، وفوج بلافاتنك الأول (29 نوفمبر 2026).')
body(doc, '3) المصدران الداخليان [1] و[11] بلا رابط عام بطبيعتهما (مدخلات تكليف وإفادة داخلية)؛ أُبقيا كما هما.')
body(doc, '4) خطوط Diodrum Arabic غير مضمّنة داخل ملفي PPTX (سلوك PowerPoint الافتراضي)؛ يلزم توافر الخطوط على أجهزة العرض داخل الهيئة — وهي خطوط الهوية المعتمدة نفسها. ملفات PDF مضمّنة الخطوط بالكامل وجاهزة للمشاركة والطباعة دون أي متطلب.')
body(doc, '5) روابط المراجع فُحصت آلياً من جهة البنية (132 رابطاً نشطاً في كل PDF)؛ ولم يُعد التحقق من صلاحية المحتوى عند كل رابط ضمن هذا الإخراج، اعتماداً على فحص المصدر بتاريخ 21 أغسطس 2026 المدون في كل مرجع.')

doc.save(os.path.join(OUT, 'Design-and-Editorial-QA-Report.docx'))
print('QA report saved')

# ---------------------------------------------------------------- XLSX
wb = openpyxl.Workbook()
hdr_fill = PatternFill('solid', fgColor=NAVY)
hdr_font = Font(name='Diodrum Arabic', color='FFFFFF', bold=True, size=10)
cell_font = Font(name='Diodrum Arabic', size=10)
thin = Border(*[Side(style='thin', color='DDDDDD')]*4)

def sheet(ws, title, headers, rows, widths, rtl=True):
    ws.title = title
    if rtl:
        ws.sheet_view.rightToLeft = True
    for j, htxt in enumerate(headers, 1):
        c = ws.cell(1, j, htxt); c.fill = hdr_fill; c.font = hdr_font
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        c.border = thin
    for i, row in enumerate(rows, 2):
        for j, val in enumerate(row, 1):
            c = ws.cell(i, j, val); c.font = cell_font
            c.alignment = Alignment(horizontal='right' if rtl else 'left',
                                    vertical='top', wrap_text=True)
            c.border = thin
    for j, wdt in enumerate(widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(j)].width = wdt

ws = wb.active
sheet(ws, 'قواعد المصطلحات', ['العربية', 'الإنجليزية المعتمدة', 'عدد الظهور في النسخة الإنجليزية النهائية', 'الحالة'],
 [['نائب الوزير', 'Vice Minister', '11', 'مطابق'],
  ['وكيل الوزارة', 'Deputy Minister', '25 (بما فيها المركّبات)', 'مطابق'],
  ['وكيل الوزارة المساعد', 'Assistant Deputy Minister', '11', 'مطابق'],
  ['مساعد الوزير', 'Assistant Minister', 'وارد ضمن المستوى الثاني حصراً', 'مطابق'],
  ['نائب المحافظ', 'Deputy Governor', '6', 'مطابق'],
  ['— محظور —', 'Undersecretary / Undersecretaries', '0', 'خالٍ'],
  ['— محظور —', 'Assistant Undersecretary', '0', 'خالٍ'],
  ['— محظور —', 'Deputy Minister بمعنى نائب الوزير', '0 (المعنى محفوظ بالفصل الاصطلاحي)', 'خالٍ'],
 ], [22, 30, 26, 12])

ws2 = wb.create_sheet()
sheet(ws2, 'التطابق الثنائي', ['الصفحة', 'الفرق المرصود آلياً', 'التفسير', 'الحكم'],
 [['1', 'ترميز تاريخ الغلاف', 'استخراج نصي مختلف لعبارة أغسطس 2026', 'اصطلاحي — لا فرق'],
  ['2 و31', '45,000/94,000 مقابل 45/94', 'الصيغة العربية المعتمدة «من 45 ألفاً إلى 94 ألف ريال»', 'اصطلاحي معتمد'],
  ['8، 10–14، 18، 21، 23–27، 33', 'أرقام أقسام/مستويات لاتينية مقابل كلمات عربية', '«Section 04» مقابل «القسم الرابع»؛ «Levels 3 and 4» مقابل «المستويين الثالث والرابع»', 'اصطلاحي — لا فرق'],
  ['23', '90 و5', '«5 days + 90 days» مقابل «خمسة أيام وتسعين يوم تطبيق» (كما في المصدر)', 'اصطلاحي — لا فرق'],
  ['32', '100,000/60,000 مقابل 100/60', '«100 ألف» و«60 ألف ريال» بالعربية', 'اصطلاحي معتمد'],
  ['بقية الصفحات', 'لا فروق رقمية', 'تطابق تام في القيم والأسعار والتواريخ والمصادر', 'مطابق'],
  ['البنية', '39 = 39 صفحة؛ الأقسام الستة + المراجع؛ ألوان المستويات موحدة', 'فحص لوحات تجميعية لكل الصفحات', 'مطابق'],
 ], [16, 30, 44, 16])

ws3 = wb.create_sheet()
sheet(ws3, 'التصحيحات التحريرية', ['#', 'النسخة', 'الموضع', 'قبل', 'بعد'],
 [['T1', 'EN', '§2.4', 'Undersecretaries', 'Deputy Ministers, Deputy Governors, and Equivalent Leaders'],
  ['T2', 'EN', '§2.5', 'Assistant Undersecretaries', 'Assistant Deputy Ministers'],
  ['T3', 'EN', '§2.3', 'merged with undersecretaries', 'merged with deputy ministers'],
  ['T4', 'EN', '§2.6/§4.5/§5.1 جداول', 'Ministry undersecretary / Assistant undersecretary', 'Deputy minister / Assistant deputy minister'],
  ['T5', 'EN', '§6.2 توصية 3', 'between minister and undersecretary', 'between the minister and the deputy minister'],
  ['T6', 'EN', '§2.3 عنوان', 'Excellency-rank Leaders', 'Vice Ministers and Excellency-Rank Leaders'],
  ['A1', 'AR', '§3.3', 'فيلق مسماة ومساءل', 'فيلق مسمّى ومُساءل'],
  ['A2', 'AR', '§4.3', 'الوكلاء ونواب المحافظون', 'وكلاء الوزارات ونواب المحافظين'],
  ['A3', 'AR', '§5.2', 'مسار الوكلاء المساعدون', 'مسار الوكلاء المساعدين'],
  ['A4', 'AR', '§5.7', 'لا يتراكم بقاءه', 'لا يتراكم عند المالك'],
  ['A5', 'AR', '§6.1(3)', 'الطلب القيادي العليا قائم', 'الطلب لدى القيادات العليا قائم'],
 ], [6, 8, 18, 34, 40])

ws4 = wb.create_sheet()
sheet(ws4, 'فحوصات الملفات', ['الفحص', 'EN', 'AR'],
 [['صحة بنية PPTX (validate.py)', 'PASSED', 'PASSED'],
  ['عدد الصفحات', '39', '39'],
  ['مقاس الصفحة', 'A4 عمودي 210×297مم', 'A4 عمودي 210×297مم'],
  ['روابط حية في PDF', '132', '132'],
  ['نص عربي مقلوب أو مفكك', '—', 'لا يوجد (فحص بصري كامل)'],
  ['جداول مقصوصة أو عناصر خارج الصفحة', 'لا يوجد', 'لا يوجد'],
  ['بقايا نصوص قالب (placeholder)', 'لا يوجد', 'لا يوجد'],
  ['المراجع [1]–[71] كاملة بالترتيب', 'نعم', 'نعم'],
 ], [40, 22, 22])

wb.save(os.path.join(OUT, 'Terminology-and-Consistency-Check.xlsx'))
print('XLSX saved')
